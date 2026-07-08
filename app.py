from werkzeug.middleware.proxy_fix import ProxyFix
import os
from flask import Flask, request, render_template, redirect, url_for, session, flash, jsonify
import pandas as pd
from sklearn.metrics.pairwise import cosine_similarity
import mysql.connector
from flask_dance.contrib.google import make_google_blueprint, google
from datetime import datetime
from werkzeug.utils import secure_filename
import PyPDF2
from docx import Document
import re
from flask_bcrypt import Bcrypt

# ======================= KONFIGURASI AWAL =======================
os.environ['OAUTHLIB_INSECURE_TRANSPORT'] = '1'
os.environ['OAUTHLIB_RELAX_TOKEN_SCOPE'] = '1'

app = Flask(__name__)
app.wsgi_app = ProxyFix(app.wsgi_app, x_proto=1, x_host=1)
app.secret_key = 'skripsi_nanad_paling_keren_se_sukabumi'  # Ganti dengan environment variable di production
\
bcrypt = Bcrypt(app)

# Konfigurasi upload file
UPLOAD_FOLDER = 'static/uploads'
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif', 'pdf', 'doc', 'docx', 'ppt', 'pptx', 'xls', 'xlsx'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs('static/uploads/artikel', exist_ok=True)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# Google OAuth blueprint
blueprint = make_google_blueprint(
    client_id="160406291861-p813l2bltsgjj5k8en11c8vna41sv33a.apps.googleusercontent.com",
    client_secret="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",
    scope=["profile", "email"],
    offline=True,
    redirect_to="beranda"
)
app.register_blueprint(blueprint, url_prefix="/login")

def get_db_connection():
    return mysql.connector.connect(
        host="localhost",
        user="root",
        password="",
        database="web_edukasi_islam",
        charset='utf8mb4',
        use_unicode=True
    )

# ======================= FUNGSI INDEKS (jalankan sekali) =======================
def create_indexes():
    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        cursor.execute("ALTER TABLE log_interaksi ADD INDEX idx_user_artikel (id_jemaah, id_artikel)")
    except: pass
    try:
        cursor.execute("ALTER TABLE rating ADD INDEX idx_user_artikel (id_jemaah, id_artikel)")
    except: pass
    try:
        cursor.execute("ALTER TABLE konsultasi ADD INDEX idx_user (id_jemaah)")
    except: pass
    try:
        cursor.execute("ALTER TABLE artikel ADD INDEX idx_kategori (id_kategori)")
    except: pass
    try:
        cursor.execute("ALTER TABLE artikel ADD INDEX idx_waktu (waktu_publish)")
    except: pass
    conn.commit()
    cursor.close()
    conn.close()
# Uncomment baris di bawah jika ingin menjalankan pembuatan indeks (jalankan sekali, lalu comment lagi)
# create_indexes()

# ======================= AUTHENTIKASI JEMAAH =======================
@app.route('/')
def landing():
    return render_template('landing.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if google.authorized:
        try:
            resp = google.get("/oauth2/v2/userinfo")
            if resp.ok:
                info = resp.json()
                email, nama, google_id = info["email"], info["name"], info["id"]
                conn = get_db_connection()
                cursor = conn.cursor(dictionary=True)
                cursor.execute("SELECT * FROM jemaah WHERE email = %s", (email,))
                user = cursor.fetchone()
                if not user:
                    # User baru daftar via Google
                    cursor.execute("INSERT INTO jemaah (google_id, nama_lengkap, email, password) VALUES (%s, %s, %s, 'google_user')", (google_id, nama, email))
                    conn.commit()
                    u_id = cursor.lastrowid
                    session['user_id'], session['nama'] = u_id, nama
                    cursor.close()
                    conn.close()
                    return redirect(url_for('pilih_minat'))
                else:
                    # User sudah ada
                    u_id = user['id_jemaah']
                    session['user_id'], session['nama'] = u_id, user['nama_lengkap']
                    cursor.close()
                    conn.close()
                    return redirect(url_for('beranda'))
        except Exception as e:
            print("Google login error:", e)
            return redirect(url_for("google.login"))

    if request.method == 'POST':
        email = request.form.get('email')
        password = request.form.get('password')
        conn = get_db_connection()
        cursor = conn.cursor(dictionary=True)
        cursor.execute("SELECT * FROM jemaah WHERE email = %s", (email,))
        user = cursor.fetchone()
        if user and bcrypt.check_password_hash(user['password'], password):
            session['user_id'], session['nama'] = user['id_jemaah'], user['nama_lengkap']
            conn.close()
            return redirect(url_for('beranda'))
        conn.close()
        return render_template('login.html', error="Email atau Password salah!")
    return render_template('login.html')

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        nama = request.form.get('nama')
        email = request.form.get('email')
        password = request.form.get('password')
        conn = get_db_connection()
        cursor = conn.cursor(dictionary=True)

        cursor.execute("SELECT id_jemaah FROM jemaah WHERE email = %s", (email,))
        existing = cursor.fetchone()
        if existing:
            conn.close()
            flash('Email sudah terdaftar. Silakan gunakan email lain.', 'danger')
            return redirect(url_for('register'))

        hashed_password = bcrypt.generate_password_hash(password).decode('utf-8')
        cursor.execute("INSERT INTO jemaah (nama_lengkap, email, password) VALUES (%s, %s, %s)", (nama, email, hashed_password))
        conn.commit()
        session['user_id'], session['nama'] = cursor.lastrowid, nama
        conn.close()
        return redirect(url_for('pilih_minat'))
    return render_template('register.html')

@app.route('/logout')
def logout():
    session.clear()
    if blueprint.token:
        del blueprint.token
    return redirect(url_for('landing'))

# ======================= LOGIN USTADZ =======================
@app.route('/login_ustadz', methods=['GET', 'POST'])
def login_ustadz():
    if request.method == 'POST':
        email = request.form.get('email')
        password = request.form.get('password')
        conn = get_db_connection()
        cursor = conn.cursor(dictionary=True)
        cursor.execute("SELECT * FROM ustadz WHERE email = %s", (email,))
        ustadz = cursor.fetchone()
        conn.close()
        if ustadz and bcrypt.check_password_hash(ustadz['password'], password):
            session['ustadz_id'] = ustadz['id_ustadz']
            session['ustadz_nama'] = ustadz['nama_ustadz']
            session['role'] = 'ustadz'
            return redirect(url_for('ustadz_dashboard', id_ustadz=ustadz['id_ustadz']))
        else:
            return render_template('login_ustadz.html', error="Email atau password salah")
    return render_template('login_ustadz.html')

@app.route('/logout_ustadz')
def logout_ustadz():
    session.pop('ustadz_id', None)
    session.pop('ustadz_nama', None)
    session.pop('role', None)
    return redirect(url_for('landing'))

# ======================= API SUBKATEGORI =======================
@app.route('/api/subkategori/<int:id_induk>')
def api_subkategori(id_induk):
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    cursor.execute("SELECT id_kategori, nama_kategori FROM kategori WHERE id_induk = %s", (id_induk,))
    results = cursor.fetchall()
    cursor.close()
    conn.close()
    return jsonify(results)

# ======================= SISTEM REKOMENDASI =======================
@app.route('/pilih_minat', methods=['GET', 'POST'])
def pilih_minat():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    if request.method == 'POST':
        minat_ids = request.form.getlist('minat')
        conn = get_db_connection()
        cursor = conn.cursor()
        for kid in minat_ids:
            cursor.execute("INSERT INTO preferensi_awal (id_jemaah, id_kategori) VALUES (%s, %s)", (session['user_id'], kid))
        conn.commit()
        conn.close()
        return redirect(url_for('beranda'))
    return render_template('pilih_minat.html')

@app.route('/beranda')
def beranda():
    if 'user_id' not in session:
        return redirect(url_for('login'))

    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    user_id = session['user_id']

    # Ambil preferensi awal user
    cursor.execute("SELECT id_kategori FROM preferensi_awal WHERE id_jemaah = %s", (user_id,))
    preferensi = [p['id_kategori'] for p in cursor.fetchall()]

    rekomendasi_final = []

    # 1. Rekomendasi berdasarkan preferensi awal
    if preferensi:
        placeholders = ','.join(['%s'] * len(preferensi))
        cursor.execute(f"""
            SELECT a.*, k.nama_kategori 
            FROM artikel a 
            JOIN kategori k ON a.id_kategori = k.id_kategori
            WHERE a.id_kategori IN ({placeholders})
            ORDER BY a.id_artikel DESC
            LIMIT 3
        """, tuple(preferensi))
        rekomendasi_final = cursor.fetchall()

    # 2. Rekomendasi berdasarkan riwayat tanya ustadz
    cursor.execute("""
        SELECT DISTINCT u.spesialisasi FROM konsultasi k 
        JOIN ustadz u ON k.id_ustadz = u.id_ustadz 
        WHERE k.id_jemaah = %s
    """, (user_id,))
    konsul_ustadz = cursor.fetchall()
    if konsul_ustadz:
        kategori_ids = []
        for u in konsul_ustadz:
            spes = u['spesialisasi'].lower()
            if 'fikih' in spes:
                kategori_ids.append(1)
            if 'akidah' in spes or 'akhlak' in spes:
                kategori_ids.append(2)
            if 'sejarah' in spes:
                kategori_ids.append(3)
        if kategori_ids:
            placeholders = ','.join(['%s'] * len(kategori_ids))
            cursor.execute(f"""
                SELECT a.*, k.nama_kategori 
                FROM artikel a 
                JOIN kategori k ON a.id_kategori = k.id_kategori
                WHERE a.id_kategori IN ({placeholders}) 
                ORDER BY a.id_artikel DESC LIMIT 2
            """, tuple(kategori_ids))
            for artikel in cursor.fetchall():
                if artikel not in rekomendasi_final:
                    rekomendasi_final.append(artikel)

    # ========== COLLABORATIVE FILTERING ==========
    print(f"[DEBUG] Memulai Collaborative Filtering untuk user_id: {user_id}")

    # 1. AMBIL RATING EKSPLISIT dari tabel rating (ambil semua rating, tidak hanya >=3)
    cursor.execute("SELECT id_jemaah, id_artikel, rating FROM rating")
    ratings_explicit = cursor.fetchall()
    print(f"[DEBUG] Rating eksplisit (bintang): {len(ratings_explicit)}")

    # 2. AMBIL RATING IMPLISIT dari log_interaksi (konversi jumlah klik)
    cursor.execute("""
        SELECT 
            l.id_jemaah, 
            l.id_artikel,
            l.jumlah_klik,
            CASE 
                WHEN l.jumlah_klik >= 10 THEN 5
                WHEN l.jumlah_klik >= 7 THEN 4
                WHEN l.jumlah_klik >= 4 THEN 3
                WHEN l.jumlah_klik >= 2 THEN 2
                WHEN l.jumlah_klik >= 1 THEN 1
                ELSE 0
            END as rating_implisit
        FROM log_interaksi l
        WHERE l.jumlah_klik > 0
    """)
    ratings_implicit = cursor.fetchall()
    print(f"[DEBUG] Jumlah data interaksi : {len(ratings_implicit)}")

    # 3. GABUNGKAN kedua sumber rating
    all_ratings = []
    for r in ratings_explicit:
        all_ratings.append({
            'id_jemaah': r['id_jemaah'],
            'id_artikel': r['id_artikel'],
            'rating': float(r['rating'])
        })
    existing_keys = {(r['id_jemaah'], r['id_artikel']) for r in all_ratings}
    for r in ratings_implicit:
        key = (r['id_jemaah'], r['id_artikel'])
        if key not in existing_keys and r['rating_implisit'] >= 1:
            all_ratings.append({
                'id_jemaah': r['id_jemaah'],
                'id_artikel': r['id_artikel'],
                'rating': float(r['rating_implisit'])
            })

    print(f"[DEBUG] Total data setelah digabung: {len(all_ratings)}")
    user_has_data = any(r['id_jemaah'] == user_id for r in all_ratings)
    print(f"[DEBUG] User {user_id} punya data? {user_has_data}")

    # 4. JALANKAN CF jika data cukup
    if len(all_ratings) >= 5 and user_has_data:
        try:
            df = pd.DataFrame(all_ratings)
            user_rating_counts = df.groupby('id_jemaah').size()
            active_users = user_rating_counts[user_rating_counts >= 2].index.tolist()
            print(f"[DEBUG] User aktif (min 2 data): {active_users}")

            if user_id in active_users and len(active_users) >= 2:
                matrix = df[df['id_jemaah'].isin(active_users)].pivot_table(
                    index='id_jemaah',
                    columns='id_artikel',
                    values='rating'
                ).fillna(0)
                print(f"[DEBUG] Matriks ukuran: {matrix.shape}")

                if matrix.shape[0] >= 2 and matrix.shape[1] >= 1:
                    similarity = cosine_similarity(matrix)
                    sim_df = pd.DataFrame(similarity, index=matrix.index, columns=matrix.index)
                    similar_users = sim_df[user_id].sort_values(ascending=False).index[1:3]
                    print(f"[DEBUG] User mirip: {list(similar_users)}")

                    cursor.execute("""
                        SELECT id_artikel FROM rating WHERE id_jemaah = %s
                        UNION
                        SELECT id_artikel FROM log_interaksi WHERE id_jemaah = %s
                    """, (user_id, user_id))
                    read_articles = [row['id_artikel'] for row in cursor.fetchall()]

                    for similar_user in similar_users:
                        if len(rekomendasi_final) >= 6:
                            break
                        if read_articles:
                            placeholders = ','.join(['%s'] * len(read_articles))
                            cursor.execute(f"""
                                SELECT DISTINCT a.*, k.nama_kategori 
                                FROM log_interaksi l
                                JOIN artikel a ON l.id_artikel = a.id_artikel
                                JOIN kategori k ON a.id_kategori = k.id_kategori
                                WHERE l.id_jemaah = %s 
                                AND l.jumlah_klik >= 1
                                AND a.id_artikel NOT IN ({placeholders})
                                ORDER BY l.jumlah_klik DESC
                                LIMIT 2
                            """, (int(similar_user),) + tuple(read_articles))
                        else:
                            cursor.execute("""
                                SELECT DISTINCT a.*, k.nama_kategori 
                                FROM log_interaksi l
                                JOIN artikel a ON l.id_artikel = a.id_artikel
                                JOIN kategori k ON a.id_kategori = k.id_kategori
                                WHERE l.id_jemaah = %s 
                                AND l.jumlah_klik >= 3
                                ORDER BY l.jumlah_klik DESC
                                LIMIT 2
                            """, (int(similar_user),))
                        recs = cursor.fetchall()
                        for rec in recs:
                            if rec not in rekomendasi_final:
                                rekomendasi_final.append(rec)
                                print(f"[DEBUG] Rekomendasi: {rec['judul_artikel']}")
        except Exception as e:
            print("[ERROR CF:]", e)
            import traceback
            traceback.print_exc()
    else:
        print("[INFO] Skip CF: Data tidak cukup")
        if len(all_ratings) < 5:
            print("  - Total data rating < 5")
        if not user_has_data:
            print("  - User belum punya aktivitas baca/rating")

    # ========== AKHIR CF ==========

    # Artikel Populer (Berdasarkan Jumlah User Unik per Artikel, tanpa RAND)
    cursor.execute("""
        SELECT a.*, k.nama_kategori, COUNT(DISTINCT l.id_jemaah) as user_count
        FROM artikel a
        JOIN kategori k ON a.id_kategori = k.id_kategori
        LEFT JOIN log_interaksi l ON a.id_artikel = l.id_artikel
        GROUP BY a.id_artikel
        ORDER BY user_count DESC, a.id_artikel DESC
        LIMIT 6
    """)
    artikel_populer = cursor.fetchall()

    icon_map = {1: 'fa-kaaba', 2: 'fa-heart', 3: 'fa-book-open'}
    for artikel in artikel_populer:
        artikel['deskripsi_singkat'] = re.sub(r'<[^>]+>', '', artikel.get('isi_materi', '') or '')[:120] + '...'
        artikel['ikon'] = icon_map.get(artikel['id_kategori'], 'fa-mosque')
    for artikel in rekomendasi_final:
        artikel['deskripsi_singkat'] = re.sub(r'<[^>]+>', '', artikel.get('isi_materi', '') or '')[:150] + '...'
        artikel['ikon'] = icon_map.get(artikel['id_kategori'], 'fa-mosque')

    # Fallback jika rekomendasi kurang dari 3
    if len(rekomendasi_final) < 3:
        cursor.execute("""
            SELECT a.*, k.nama_kategori
            FROM artikel a
            JOIN kategori k ON a.id_kategori = k.id_kategori
            ORDER BY a.id_artikel DESC
            LIMIT 6
        """)
        populer = cursor.fetchall()
        existing_ids = [r['id_artikel'] for r in rekomendasi_final]
        for artikel in populer:
            if artikel['id_artikel'] not in existing_ids:
                rekomendasi_final.append(artikel)
            if len(rekomendasi_final) >= 6:
                break

    cursor.close()
    conn.close()
    return render_template('index.html',
                          rekomendasi_cf=rekomendasi_final[:6],
                          artikel_populer=artikel_populer,
                          user_preferensi=preferensi)

# ======================= E-COUNSELING =======================
@app.route('/ustadz')
def ustadz_list():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    cursor.execute("SELECT * FROM ustadz")
    res = cursor.fetchall()
    cursor.close()
    conn.close()
    return render_template('ustadz_list.html', ustadz=res)

@app.route('/ustadz/<int:id_ustadz>', methods=['GET', 'POST'])
def ustadz_detail(id_ustadz):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    if request.method == 'POST':
        tanya = request.form.get('pertanyaan')
        cursor.execute("INSERT INTO konsultasi (id_jemaah, id_ustadz, pertanyaan, waktu_kirim) VALUES (%s, %s, %s, NOW())",
                       (session['user_id'], id_ustadz, tanya))
        conn.commit()
        conn.close()
        return redirect(url_for('profil'))
    cursor.execute("SELECT * FROM ustadz WHERE id_ustadz = %s", (id_ustadz,))
    u = cursor.fetchone()
    cursor.close()
    conn.close()
    return render_template('ustadz_detail.html', ustadz=u)

@app.route('/ustadz/dashboard/<int:id_ustadz>')
def ustadz_dashboard(id_ustadz):
    if 'ustadz_id' not in session or session['ustadz_id'] != id_ustadz:
        return redirect(url_for('login_ustadz'))
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    cursor.execute("SELECT * FROM ustadz WHERE id_ustadz = %s", (id_ustadz,))
    u = cursor.fetchone()
    cursor.execute("""
        SELECT k.id_konsultasi, k.id_jemaah, k.pertanyaan, k.waktu_kirim, j.nama_lengkap as nama_jemaah
        FROM konsultasi k
        JOIN jemaah j ON k.id_jemaah = j.id_jemaah
        WHERE k.id_ustadz = %s
        ORDER BY k.waktu_kirim ASC
    """, (id_ustadz,))
    rows = cursor.fetchall()
    jawaban_map = {}
    if rows:
        ids = [str(row['id_konsultasi']) for row in rows]
        placeholders = ','.join(['%s'] * len(ids))
        cursor.execute(f"""
            SELECT id_konsultasi, jawaban, created_at, ustadz_id
            FROM konsultasi_jawaban
            WHERE id_konsultasi IN ({placeholders})
            ORDER BY created_at ASC
        """, ids)
        jawaban_rows = cursor.fetchall()
        for jr in jawaban_rows:
            kid = jr['id_konsultasi']
            if kid not in jawaban_map:
                jawaban_map[kid] = []
            jawaban_map[kid].append({
                'jawaban': jr['jawaban'],
                'created_at': jr['created_at'],
                'ustadz_id': jr['ustadz_id']
            })
    cursor.close()
    conn.close()
    rooms = {}
    for row in rows:
        jid = row['id_jemaah']
        if jid not in rooms:
            rooms[jid] = {'nama_jemaah': row['nama_jemaah'], 'messages': []}
        rooms[jid]['messages'].append({
            'id_konsultasi': row['id_konsultasi'],
            'pertanyaan': row['pertanyaan'],
            'waktu_kirim': row['waktu_kirim'],
            'jawaban_list': jawaban_map.get(row['id_konsultasi'], [])
        })
    now = datetime.now().strftime("%d %B %Y")
    return render_template('ustadz_dashboard.html', ustadz=u, rooms=rooms, now=now)

@app.route('/balas_konsultasi', methods=['POST'])
def balas_konsultasi():
    if 'ustadz_id' not in session:
        return redirect(url_for('login_ustadz'))
    id_kon = request.form.get('id_konsultasi')
    jawaban = request.form.get('jawaban')
    id_u = request.form.get('id_ustadz')
    if not id_kon or not jawaban:
        return redirect(url_for('ustadz_dashboard', id_ustadz=id_u))
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("INSERT INTO konsultasi_jawaban (id_konsultasi, jawaban, ustadz_id, created_at) VALUES (%s, %s, %s, NOW())",
                   (id_kon, jawaban, session['ustadz_id']))
    conn.commit()
    cursor.close()
    conn.close()
    return redirect(url_for('ustadz_dashboard', id_ustadz=id_u))

@app.route('/kirim_pertanyaan_lanjutan', methods=['POST'])
def kirim_pertanyaan_lanjutan():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    id_ustadz = request.form.get('id_ustadz')
    pertanyaan = request.form.get('pertanyaan')
    if pertanyaan and id_ustadz:
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("INSERT INTO konsultasi (id_jemaah, id_ustadz, pertanyaan, waktu_kirim) VALUES (%s, %s, %s, NOW())",
                       (session['user_id'], id_ustadz, pertanyaan))
        conn.commit()
        conn.close()
    return redirect(url_for('profil'))

@app.route('/hapus_room/<int:id_ustadz>')
def hapus_room(id_ustadz):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("""
        DELETE kj FROM konsultasi_jawaban kj
        INNER JOIN konsultasi k ON kj.id_konsultasi = k.id_konsultasi
        WHERE k.id_jemaah = %s AND k.id_ustadz = %s
    """, (session['user_id'], id_ustadz))
    cursor.execute("DELETE FROM konsultasi WHERE id_jemaah = %s AND id_ustadz = %s", (session['user_id'], id_ustadz))
    conn.commit()
    cursor.close()
    conn.close()
    return redirect(url_for('profil'))

@app.route('/hapus_konsultasi_ustadz/<int:id_jemaah>')
def hapus_konsultasi_ustadz(id_jemaah):
    if 'ustadz_id' not in session:
        return redirect(url_for('login_ustadz'))
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("""
        DELETE kj FROM konsultasi_jawaban kj
        INNER JOIN konsultasi k ON kj.id_konsultasi = k.id_konsultasi
        WHERE k.id_ustadz = %s AND k.id_jemaah = %s
    """, (session['ustadz_id'], id_jemaah))
    cursor.execute("DELETE FROM konsultasi WHERE id_ustadz = %s AND id_jemaah = %s", (session['ustadz_id'], id_jemaah))
    conn.commit()
    cursor.close()
    conn.close()
    return redirect(url_for('ustadz_dashboard', id_ustadz=session['ustadz_id']))

# ======================= ARTIKEL & PROFIL =======================
@app.route('/baca/<int:id_artikel>')
def baca_artikel(id_artikel):
    if 'user_id' not in session:
        return redirect(url_for('login'))

    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)

    # LOG INTERAKSI
    user_id = session['user_id']
    cursor.execute("SELECT * FROM log_interaksi WHERE id_jemaah = %s AND id_artikel = %s", (user_id, id_artikel))
    existing = cursor.fetchone()
    if existing:
        cursor.execute("UPDATE log_interaksi SET jumlah_klik = jumlah_klik + 1, terakhir_dibaca = NOW() WHERE id_jemaah = %s AND id_artikel = %s", (user_id, id_artikel))
    else:
        cursor.execute("INSERT INTO log_interaksi (id_jemaah, id_artikel, jumlah_klik, terakhir_dibaca) VALUES (%s, %s, 1, NOW())", (user_id, id_artikel))
    conn.commit()

    # AMBIL DATA ARTIKEL
    cursor.execute("SELECT a.*, k.nama_kategori FROM artikel a JOIN kategori k ON a.id_kategori = k.id_kategori WHERE a.id_artikel = %s", (id_artikel,))
    art = cursor.fetchone()
    cursor.execute("SELECT rating FROM rating WHERE id_jemaah = %s AND id_artikel = %s", (session['user_id'], id_artikel))
    rtg = cursor.fetchone()

    # REKOMENDASI ARTIKEL TERKAIT (item-based: artikel yang sering dibaca bersama)
    cursor.execute("""
        SELECT l2.id_artikel, a.judul_artikel, a.id_kategori, k.nama_kategori, COUNT(*) as bersama
        FROM log_interaksi l1
        JOIN log_interaksi l2 ON l1.id_jemaah = l2.id_jemaah AND l1.id_artikel != l2.id_artikel
        JOIN artikel a ON l2.id_artikel = a.id_artikel
        JOIN kategori k ON a.id_kategori = k.id_kategori
        WHERE l1.id_artikel = %s
        GROUP BY l2.id_artikel
        ORDER BY bersama DESC
        LIMIT 3
    """, (id_artikel,))
    rekom_artikel = cursor.fetchall()
    if not rekom_artikel:
        cursor.execute("SELECT a.*, k.nama_kategori FROM artikel a JOIN kategori k ON a.id_kategori = k.id_kategori ORDER BY a.id_artikel DESC LIMIT 3")
        rekom_artikel = cursor.fetchall()

    cursor.close()
    conn.close()
    return render_template('baca_artikel.html', artikel=art, user_rating=rtg, artikel_rekomendasi=rekom_artikel)

@app.route('/submit_rating', methods=['POST'])
def submit_rating():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    id_art = request.form.get('id_artikel')
    r_val = request.form.get('rating')
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("INSERT INTO rating (id_jemaah, id_artikel, rating) VALUES (%s, %s, %s) ON DUPLICATE KEY UPDATE rating = %s", (session['user_id'], id_art, r_val, r_val))
    conn.commit()
    conn.close()
    return redirect(url_for('baca_artikel', id_artikel=id_art))

@app.route('/profil')
def profil():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    cursor.execute("SELECT * FROM jemaah WHERE id_jemaah = %s", (session['user_id'],))
    usr = cursor.fetchone()
    cursor.execute("""
        SELECT 
            a.judul_artikel, 
            a.id_artikel,
            COALESCE(r.rating, 0) as rating,
            l.jumlah_klik,
            l.terakhir_dibaca as waktu_baca
        FROM log_interaksi l
        JOIN artikel a ON l.id_artikel = a.id_artikel
        LEFT JOIN rating r ON r.id_jemaah = l.id_jemaah AND r.id_artikel = l.id_artikel
        WHERE l.id_jemaah = %s
        ORDER BY l.terakhir_dibaca DESC
    """, (session['user_id'],))
    riwayat = cursor.fetchall()
    cursor.execute("""
        SELECT k.id_konsultasi, k.id_ustadz, k.pertanyaan, k.waktu_kirim, u.nama_ustadz
        FROM konsultasi k
        JOIN ustadz u ON k.id_ustadz = u.id_ustadz
        WHERE k.id_jemaah = %s
        ORDER BY k.waktu_kirim ASC
    """, (session['user_id'],))
    konsul_rows = cursor.fetchall()
    jawaban_map = {}
    if konsul_rows:
        ids = [str(row['id_konsultasi']) for row in konsul_rows]
        placeholders = ','.join(['%s'] * len(ids))
        cursor.execute(f"""
            SELECT id_konsultasi, jawaban, created_at, ustadz_id
            FROM konsultasi_jawaban
            WHERE id_konsultasi IN ({placeholders})
            ORDER BY created_at ASC
        """, ids)
        jawaban_rows = cursor.fetchall()
        for jr in jawaban_rows:
            kid = jr['id_konsultasi']
            if kid not in jawaban_map:
                jawaban_map[kid] = []
            jawaban_map[kid].append({
                'jawaban': jr['jawaban'],
                'created_at': jr['created_at'],
                'ustadz_id': jr['ustadz_id']
            })
    cursor.close()
    conn.close()
    rooms = {}
    for row in konsul_rows:
        uid = row['id_ustadz']
        if uid not in rooms:
            rooms[uid] = {
                'nama_ustadz': row['nama_ustadz'],
                'messages': []
            }
        rooms[uid]['messages'].append({
            'id_konsultasi': row['id_konsultasi'],
            'pertanyaan': row['pertanyaan'],
            'waktu_kirim': row['waktu_kirim'],
            'jawaban_list': jawaban_map.get(row['id_konsultasi'], [])
        })
    return render_template('profil.html', user=usr, riwayat=riwayat, rooms=rooms)

# ======================= HALAMAN KATEGORI =======================
@app.route('/kategori/<int:id_kategori>')
def kategori(id_kategori):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    cursor.execute("SELECT * FROM kategori WHERE id_kategori = %s", (id_kategori,))
    kat = cursor.fetchone()
    if not kat:
        cursor.close()
        conn.close()
        return redirect(url_for('beranda'))
    cursor.execute("""
        SELECT a.*, k.nama_kategori 
        FROM artikel a 
        JOIN kategori k ON a.id_kategori = k.id_kategori
        WHERE a.id_kategori = %s
        ORDER BY a.id_artikel DESC
    """, (id_kategori,))
    artikel_list = cursor.fetchall()
    for artikel in artikel_list:
        if artikel.get('isi_materi'):
            artikel['preview'] = re.sub(r'<[^>]+>', '', artikel['isi_materi'])[:200] + '...'
        else:
            artikel['preview'] = 'Tidak ada konten'
        icon_map = {1: 'fa-kaaba', 2: 'fa-heart', 3: 'fa-book-open'}
        artikel['ikon'] = icon_map.get(artikel['id_kategori'], 'fa-mosque')
    cursor.close()
    conn.close()
    return render_template('kategori.html', kategori=kat, artikel_list=artikel_list)

# ======================= FITUR PENCARIAN =======================
@app.route('/cari')
def cari():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    keyword = request.args.get('q', '').strip()
    hasil = []
    if keyword:
        conn = get_db_connection()
        cursor = conn.cursor(dictionary=True)
        cursor.execute("""
            SELECT a.*, k.nama_kategori 
            FROM artikel a 
            JOIN kategori k ON a.id_kategori = k.id_kategori
            WHERE a.judul_artikel LIKE %s 
               OR a.isi_materi LIKE %s 
               OR a.sumber LIKE %s
            ORDER BY 
                (CASE WHEN a.judul_artikel LIKE %s THEN 1 ELSE 0 END) DESC,
                a.id_artikel DESC
        """, (f'%{keyword}%', f'%{keyword}%', f'%{keyword}%', f'%{keyword}%'))
        hasil = cursor.fetchall()
        for item in hasil:
            if item.get('isi_materi'):
                preview = re.sub(r'<[^>]+>', '', item['isi_materi'])
                if len(preview) > 200:
                    preview = preview[:200] + '...'
                item['preview'] = preview
            else:
                item['preview'] = 'Tidak ada konten'
            icon_map = {1: 'fa-kaaba', 2: 'fa-heart', 3: 'fa-book-open'}
            item['ikon'] = icon_map.get(item['id_kategori'], 'fa-mosque')
        cursor.close()
        conn.close()
    return render_template('cari.html', keyword=keyword, hasil=hasil, total=len(hasil))

# ======================= ADMIN PANEL =======================
@app.route('/admin/login', methods=['GET', 'POST'])
def admin_login():
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        conn = get_db_connection()
        cursor = conn.cursor(dictionary=True)
        cursor.execute("SELECT * FROM admin WHERE username = %s", (username,))
        admin = cursor.fetchone()
        conn.close()
        if admin and bcrypt.check_password_hash(admin['password'], password):
            session['admin_id'] = admin['id_admin']
            session['admin_username'] = admin['username']
            return redirect(url_for('admin_dashboard'))
        else:
            return render_template('admin_login.html', error="Username atau password salah")
    return render_template('admin_login.html')

@app.route('/admin/ustadz/edit/<int:id_ustadz>', methods=['GET', 'POST'])
def admin_edit_ustadz(id_ustadz):
    if 'admin_id' not in session:
        return redirect(url_for('admin_login'))
    
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    
    if request.method == 'POST':
        nama = request.form.get('nama_ustadz')
        spesialisasi = request.form.get('spesialisasi')
        pendidikan = request.form.get('pendidikan')
        bio = request.form.get('bio')
        email = request.form.get('email')
        password_baru = request.form.get('password')
        
        if password_baru and password_baru.strip() != '':
            hashed = bcrypt.generate_password_hash(password_baru).decode('utf-8')
            cursor.execute("""
                UPDATE ustadz 
                SET nama_ustadz=%s, spesialisasi=%s, pendidikan=%s, bio=%s, email=%s, password=%s
                WHERE id_ustadz=%s
            """, (nama, spesialisasi, pendidikan, bio, email, hashed, id_ustadz))
        else:
            cursor.execute("""
                UPDATE ustadz 
                SET nama_ustadz=%s, spesialisasi=%s, pendidikan=%s, bio=%s, email=%s
                WHERE id_ustadz=%s
            """, (nama, spesialisasi, pendidikan, bio, email, id_ustadz))
        
        conn.commit()
        conn.close()
        flash('Data ustadz berhasil diupdate', 'success')
        return redirect(url_for('admin_dashboard'))
    
    # GET: tampilkan form edit
    cursor.execute("SELECT * FROM ustadz WHERE id_ustadz = %s", (id_ustadz,))
    ustadz = cursor.fetchone()
    conn.close()
    return render_template('admin_form_ustadz_edit.html', ustadz=ustadz)

@app.route('/admin/logout')
def admin_logout():
    session.pop('admin_id', None)
    session.pop('admin_username', None)
    return redirect(url_for('admin_login'))

@app.route('/admin/dashboard')
def admin_dashboard():
    if 'admin_id' not in session:
        return redirect(url_for('admin_login'))
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    cursor.execute("""
        SELECT a.*, k.nama_kategori 
        FROM artikel a 
        JOIN kategori k ON a.id_kategori = k.id_kategori
        ORDER BY a.id_artikel DESC
    """)
    artikel = cursor.fetchall()
    cursor.execute("SELECT * FROM ustadz")
    ustadz = cursor.fetchall()
    conn.close()
    now = datetime.now().strftime("%d %B %Y")
    return render_template('admin_dashboard.html', artikel=artikel, ustadz=ustadz, now=now)

def extract_text_from_pdf(filepath):
    text = ""
    try:
        with open(filepath, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"Error membaca PDF: {e}")
    return text

def extract_text_from_docx(filepath):
    text = ""
    try:
        doc = Document(filepath)
        for paragraph in doc.paragraphs:
            if paragraph.text:
                text += paragraph.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        text += cell.text + "\n"
    except Exception as e:
        print(f"Error membaca Word: {e}")
    return text

@app.route('/admin/artikel/tambah', methods=['GET', 'POST'])
def admin_tambah_artikel():
    if 'admin_id' not in session:
        return redirect(url_for('admin_login'))
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    cursor.execute("SHOW COLUMNS FROM kategori")
    columns = [col['Field'] for col in cursor.fetchall()]
    if 'parent_id' in columns:
        cursor.execute("SELECT * FROM kategori WHERE parent_id IS NULL OR parent_id = 0 ORDER BY id_kategori")
    elif 'id_induk' in columns:
        cursor.execute("SELECT * FROM kategori WHERE id_induk IS NULL OR id_induk = 0 ORDER BY id_kategori")
    else:
        cursor.execute("SELECT * FROM kategori ORDER BY id_kategori")
    daftar_kategori = cursor.fetchall()
    cursor.close()
    conn.close()
    if request.method == 'POST':
        judul = request.form.get('judul')
        isi_materi = request.form.get('isi_materi')
        kategori = request.form.get('id_kategori')
        sumber = request.form.get('sumber')
        gambar = request.form.get('link_gambar')
        file_path = None
        file_type = None
        extracted_text = ""
        if 'file_attachment' in request.files:
            file = request.files['file_attachment']
            if file and file.filename != '' and allowed_file(file.filename):
                filename = secure_filename(file.filename)
                name_parts = filename.rsplit('.', 1)
                extension = name_parts[1].lower()
                unique_filename = f"{name_parts[0]}_{datetime.now().strftime('%Y%m%d%H%M%S')}.{extension}"
                filepath = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
                file.save(filepath)
                file_path = f'/static/uploads/{unique_filename}'
                file_type = extension
                if extension == 'pdf':
                    extracted_text = extract_text_from_pdf(filepath)
                    if extracted_text:
                        flash(f'Berhasil membaca isi PDF ({len(extracted_text)} karakter)', 'success')
                    else:
                        flash('Gagal membaca isi PDF. Silakan isi manual.', 'warning')
                elif extension in ['doc', 'docx']:
                    extracted_text = extract_text_from_docx(filepath)
                    if extracted_text:
                        flash(f'Berhasil membaca isi Word ({len(extracted_text)} karakter)', 'success')
                    else:
                        flash('Gagal membaca isi Word. Silakan isi manual.', 'warning')
                if (not isi_materi or isi_materi.strip() == '') and extracted_text:
                    if len(extracted_text) > 10000:
                        extracted_text = extracted_text[:10000] + "\n\n... (teks dipotong karena terlalu panjang)"
                    isi_materi = extracted_text
        if not isi_materi or isi_materi.strip() == '':
            flash('Isi materi tidak boleh kosong. Silakan tulis manual atau unggah file PDF/Word yang valid.', 'danger')
            return render_template('admin_form_artikel.html', kategori_list=daftar_kategori)
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("""
            INSERT INTO artikel (judul_artikel, isi_materi, id_kategori, sumber, link_gambar, file_attachments, file_type, waktu_publish)
            VALUES (%s, %s, %s, %s, %s, %s, %s, NOW())
        """, (judul, isi_materi, kategori, sumber, gambar, file_path, file_type))
        conn.commit()
        conn.close()
        flash('Artikel berhasil ditambahkan!', 'success')
        return redirect(url_for('admin_dashboard'))
    return render_template('admin_form_artikel.html', kategori_list=daftar_kategori)

@app.route('/admin/artikel/edit/<int:id_artikel>', methods=['GET', 'POST'])
def admin_edit_artikel(id_artikel):
    if 'admin_id' not in session:
        return redirect(url_for('admin_login'))
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    cursor.execute("SHOW COLUMNS FROM kategori")
    columns = [col['Field'] for col in cursor.fetchall()]
    if 'parent_id' in columns:
        cursor.execute("SELECT * FROM kategori WHERE parent_id IS NULL OR parent_id = 0 ORDER BY id_kategori")
    elif 'id_induk' in columns:
        cursor.execute("SELECT * FROM kategori WHERE id_induk IS NULL OR id_induk = 0 ORDER BY id_kategori")
    else:
        cursor.execute("SELECT * FROM kategori ORDER BY id_kategori")
    daftar_kategori = cursor.fetchall()
    if request.method == 'POST':
        judul = request.form.get('judul')
        isi_materi = request.form.get('isi_materi')
        kategori = request.form.get('id_kategori')
        sumber = request.form.get('sumber')
        gambar = request.form.get('link_gambar')
        file_path = None
        file_type = None
        extracted_text = ""
        if 'file_attachment' in request.files:
            file = request.files['file_attachment']
            if file and file.filename != '' and allowed_file(file.filename):
                cursor.execute("SELECT file_attachments FROM artikel WHERE id_artikel = %s", (id_artikel,))
                old_file = cursor.fetchone()
                if old_file and old_file.get('file_attachments'):
                    old_file_path = old_file['file_attachments'].lstrip('/')
                    if os.path.exists(old_file_path):
                        os.remove(old_file_path)
                filename = secure_filename(file.filename)
                name_parts = filename.rsplit('.', 1)
                extension = name_parts[1].lower()
                unique_filename = f"{name_parts[0]}_{datetime.now().strftime('%Y%m%d%H%M%S')}.{extension}"
                filepath = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
                file.save(filepath)
                file_path = f'/static/uploads/{unique_filename}'
                file_type = extension
                if extension == 'pdf':
                    extracted_text = extract_text_from_pdf(filepath)
                elif extension in ['doc', 'docx']:
                    extracted_text = extract_text_from_docx(filepath)
                if (not isi_materi or isi_materi.strip() == '') and extracted_text:
                    if len(extracted_text) > 10000:
                        extracted_text = extracted_text[:10000] + "\n\n... (teks dipotong)"
                    isi_materi = extracted_text
                    flash(f'Isi materi otomatis diisi dari file {extension.upper()}', 'success')
                cursor.execute("UPDATE artikel SET file_attachments=%s, file_type=%s WHERE id_artikel=%s", (file_path, file_type, id_artikel))
        if not isi_materi or isi_materi.strip() == '':
            flash('Isi materi tidak boleh kosong!', 'danger')
            return redirect(url_for('admin_edit_artikel', id_artikel=id_artikel))
        cursor.execute("""
            UPDATE artikel SET judul_artikel=%s, isi_materi=%s, id_kategori=%s, sumber=%s, link_gambar=%s
            WHERE id_artikel=%s
        """, (judul, isi_materi, kategori, sumber, gambar, id_artikel))
        conn.commit()
        conn.close()
        flash('Artikel berhasil diupdate!', 'success')
        return redirect(url_for('admin_dashboard'))
    cursor.execute("SELECT * FROM artikel WHERE id_artikel = %s", (id_artikel,))
    artikel = cursor.fetchone()
    conn.close()
    return render_template('admin_form_artikel.html', artikel=artikel, kategori_list=daftar_kategori)

@app.route('/admin/artikel/hapus/<int:id_artikel>')
def admin_hapus_artikel(id_artikel):
    if 'admin_id' not in session:
        return redirect(url_for('admin_login'))
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    cursor.execute("SELECT file_attachments FROM artikel WHERE id_artikel = %s", (id_artikel,))
    artikel = cursor.fetchone()
    if artikel and artikel.get('file_attachments'):
        file_path = artikel['file_attachments'].lstrip('/')
        if os.path.exists(file_path):
            os.remove(file_path)
    cursor.execute("DELETE FROM artikel WHERE id_artikel = %s", (id_artikel,))
    conn.commit()
    conn.close()
    return redirect(url_for('admin_dashboard'))

@app.route('/admin/artikel/hapus_file/<int:id_artikel>')
def admin_hapus_file(id_artikel):
    if 'admin_id' not in session:
        return redirect(url_for('admin_login'))
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    cursor.execute("SELECT file_attachments FROM artikel WHERE id_artikel = %s", (id_artikel,))
    artikel = cursor.fetchone()
    if artikel and artikel.get('file_attachments'):
        file_path = artikel['file_attachments'].lstrip('/')
        if os.path.exists(file_path):
            os.remove(file_path)
        cursor.execute("UPDATE artikel SET file_attachments = NULL, file_type = NULL WHERE id_artikel = %s", (id_artikel,))
        conn.commit()
    conn.close()
    return redirect(url_for('admin_edit_artikel', id_artikel=id_artikel))

@app.route('/admin/ustadz/tambah', methods=['GET', 'POST'])
def admin_tambah_ustadz():
    if 'admin_id' not in session:
        return redirect(url_for('admin_login'))
    if request.method == 'POST':
        nama = request.form.get('nama_ustadz')
        spesialisasi = request.form.get('spesialisasi')
        pendidikan = request.form.get('pendidikan')
        bio = request.form.get('bio')
        email = request.form.get('email')
        password = request.form.get('password')
        hashed = bcrypt.generate_password_hash(password).decode('utf-8')
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("""
            INSERT INTO ustadz (nama_ustadz, spesialisasi, pendidikan, bio, email, password)
            VALUES (%s, %s, %s, %s, %s, %s)
        """, (nama, spesialisasi, pendidikan, bio, email, hashed))
        conn.commit()
        conn.close()
        return redirect(url_for('admin_dashboard'))
    return render_template('admin_form_ustadz.html')

@app.route('/admin/ustadz/hapus/<int:id_ustadz>')
def admin_hapus_ustadz(id_ustadz):
    if 'admin_id' not in session:
        return redirect(url_for('admin_login'))
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("DELETE FROM ustadz WHERE id_ustadz = %s", (id_ustadz,))
    conn.commit()
    conn.close()
    return redirect(url_for('admin_dashboard'))

# ======================= UPLOAD FOTO PROFIL =======================
@app.route('/upload_foto', methods=['POST'])
def upload_foto():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    if 'foto' not in request.files:
        flash('Tidak ada file yang dipilih', 'danger')
        return redirect(url_for('profil'))
    file = request.files['foto']
    if file.filename == '':
        flash('File kosong', 'danger')
        return redirect(url_for('profil'))
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        ext = filename.rsplit('.', 1)[1].lower()
        new_filename = f"user_{session['user_id']}.{ext}"
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], new_filename)
        file.save(filepath)
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("UPDATE jemaah SET foto_profil = %s WHERE id_jemaah = %s", (f'/static/uploads/{new_filename}', session['user_id']))
        conn.commit()
        cursor.close()
        conn.close()
        flash('Foto profil berhasil diupdate', 'success')
    else:
        flash('Format file tidak diizinkan (png, jpg, jpeg, gif)', 'danger')
    return redirect(url_for('profil'))

@app.route('/edit_profil', methods=['POST'])
def edit_profil():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    bio = request.form.get('bio')
    no_telepon = request.form.get('no_telepon')
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("UPDATE jemaah SET bio = %s, no_telepon = %s WHERE id_jemaah = %s", (bio, no_telepon, session['user_id']))
    conn.commit()
    conn.close()
    flash('Profil berhasil diupdate', 'success')
    return redirect(url_for('profil'))

@app.route('/hapus_foto')
def hapus_foto():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("UPDATE jemaah SET foto_profil = NULL WHERE id_jemaah = %s", (session['user_id'],))
    conn.commit()
    cursor.close()
    conn.close()
    flash('Foto profil berhasil dihapus', 'success')
    return redirect(url_for('profil'))

# ======================= MAIN =======================
if __name__ == '__main__':
    app.run(debug=True)
